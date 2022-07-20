#!/usr/bin/env python3
from docx import Document
from docx.table import _Cell
from docx.shared import Mm, Cm, Pt
from docx.enum.text import WD_LINE_SPACING, WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement,qn
from progress.bar import FillingSquaresBar
import io

class WORDDOCX(object):
	def __init__(self, template_filename):
		self.doc = Document(template_filename)

	def set_repeat_table_header(self, row):
		""" set repeat table row on every new page
		"""
		tr = row._tr
		trPr = tr.get_or_add_trPr()
		tblHeader = OxmlElement('w:tblHeader')
		tblHeader.set(qn('w:val'), "true")
		trPr.append(tblHeader)
		return row

	def set_cell_margins(self, cell: _Cell, **kwargs):
		"""
		cell:  actual cell instance you want to modify
		usage:
			set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
		provided values are in twentieths of a point (1/1440 of an inch).
		read more here: http://officeopenxml.com/WPtableCellMargins.php
		"""
		tc = cell._tc
		tcPr = tc.get_or_add_tcPr()
		tcMar = OxmlElement('w:tcMar')
	 
		for m in [
			"top",
			"start",
			"bottom",
			"end",
		]:
			if m in kwargs:
				node = OxmlElement("w:{}".format(m))
				node.set(qn('w:w'), str(kwargs.get(m)))
				node.set(qn('w:type'), 'dxa')
				tcMar.append(node)
		tcPr.append(tcMar)
		return cell

	def create_nessus_docx(self, report_filename, vulns):
		doc = self.doc
		table = doc.add_table(1, 6, 'Table Grid')
		table.autofit = False
		head_cells = table.rows[0].cells
		self.set_repeat_table_header(table.rows[0])

		for i, item in enumerate(['№', 'CVE / Название уязвимости', 'Уровень опасности уязвимости', 'Хост / Приложение / Раздел', 'Описание уязвимости и воздействия', 'Рекомендации']):
			p = head_cells[i].paragraphs[0]
			fmt = p.paragraph_format
			fmt.first_line_indent = Mm(0)
			fmt.left_indent = Mm(0)
			fmt.right_indent = Mm(0)
			fmt.space_before = Mm(0)
			fmt.space_after = Mm(0)
			fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
			txt = p.add_run(item)
			txt.bold = True
			txt.font.size = Pt(10)
			txt.font.name = 'Times New Roman'

		items = []
		bar = FillingSquaresBar('Status', max = len(vulns), suffix='%(index)d/%(max)d vulns [%(elapsed)ds / %(eta)ds]')
		for i, (k, v) in enumerate(vulns.items(), start = 1):
			bar.next()
			plugName = v['pluginName']
			sev = v['severity']
			if sev == '0':
				continue
				# severity = 'Инфо'
			elif sev == '1':
				severity = 'Низкий'
			elif sev == '2':
				severity = 'Средний'
			elif sev == '3':
				severity = 'Высокий'
			elif sev == '4':
				severity = 'Критический'
			else:
				severity = sev
			descr = v['description']
			solution = v['solution']
			if 'see_also' in v:
				see_also = v['see_also']
			else:
				see_also = ''
			if 'cve' in v:
				cve = v['cve']
				if type(cve) == list:
					cves = " ({0})".format(', '.join(cve))
				else:
					cves = " ({0})".format(cve)
			else:
				cves = ''
			targets = ', '.join(v['target'])
			# hostname = plugID['host_fqdn']
			if see_also != '':
				recom = "{0} Подробнее: {1}".format(solution, see_also)
			else:
				recom = solution
			vulnName = "{0}{1}".format(plugName, cves)
			item = [i, vulnName, severity, targets, descr, recom]
			items.append(item)
		bar.finish()

		for row in items:
			cells = table.add_row().cells
			for i, item in enumerate(row):
				cells[i].text = str(item)
				p = cells[i].paragraphs[0]
				fmt = p.paragraph_format
				fmt.first_line_indent = Mm(0)
				fmt.left_indent = Mm(0)
				fmt.right_indent = Mm(0)
				fmt.space_before = Mm(0)
				fmt.space_after = Mm(0)
				fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
				txt = p.runs[0]
				txt.font.size = Pt(10)
				txt.font.name = 'Times New Roman'

		# width in Cm
		for i, item in enumerate(['0.75', '3.75', '2.25', '4.25', '7.75', '7']):
			for cell in table.columns[i].cells:
				# 1 Cm = 360 000 EMU (English Metric Unit)
				w = int(float(item)*360000)
				cell.width = w
				# cell = self.set_cell_margins(cell, start=50, end=50)
		doc.save(report_filename)
		print("\nDocx file '{0}' has been created".format(report_filename))
